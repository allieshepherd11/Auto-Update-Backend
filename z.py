from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd

# Sample DataFrame
data = {'Column1': [1, 2, 3], 'Column2': ['A', 'B', 'C'], 'Column3': [4.5, 5.5, 6.5]}
df = pd.DataFrame(data)

# Create a Document
doc = Document()

# Add a table with the same number of columns as the DataFrame
table = doc.add_table(rows=1, cols=len(df.columns))

# Function to make text bold and underlined
def make_bold_underline(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.underline = True

# Function to remove table borders
def remove_table_borders(tbl):
    tbl_pr = tbl._element.find(qn('w:tblPr'))  # Get <w:tblPr>
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl._element.insert(0, tbl_pr)  # Insert at the beginning

    # Create or replace <w:tblBorders> to remove borders
    tbl_borders = tbl_pr.find(qn('w:tblBorders'))
    if tbl_borders is None:
        tbl_borders = OxmlElement('w:tblBorders')
        tbl_pr.append(tbl_borders)

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')  # 'nil' removes the border
        tbl_borders.append(border)

# Set table header (First Row - Bold & Underlined)
hdr_cells = table.rows[0].cells
for col_idx, column_name in enumerate(df.columns):
    hdr_cells[col_idx].text = column_name
    make_bold_underline(hdr_cells[col_idx])  # Apply bold and underline

# Populate table with data
rows = []
for _, row in df.iterrows():
    row_cells = table.add_row().cells
    for col_idx, value in enumerate(row):
        row_cells[col_idx].text = str(value)
    rows.append(row_cells)

# Make last row bold & underlined
last_row_cells = rows[-1]  # Get the last row added
for cell in last_row_cells:
    make_bold_underline(cell)

# Remove table borders
remove_table_borders(table)

# Save the document
doc.save("output.docx")

print("Word document saved as 'output.docx' with first and last row bolded & underlined, and no table borders.")
