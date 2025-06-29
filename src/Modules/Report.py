import pandas as pd
import numpy as np
from fpdf import FPDF
import datetime
import json
import time
import unicodedata
import warnings
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd
from docx import Document
import datetime as dt
import os
import glob


warnings.filterwarnings('ignore')


class ProdReport():
    def __init__(self, field:str,title:str,clientName=None,wells=None,day=None):
        self.field = field
        self.wells = wells
        self.clientName = clientName
        self.title = title.title().replace(' Re ',' RE ')
        self.dtype = {'Oil (BBLS)': float, 'Water (BBLS)': float, 'Gas (MCF)': float, 'TP': float, 'CP': float, 'Comments': str}
        self.df = pd.read_csv(f'data/prod/{self.field}/data.csv', dtype=self.dtype)
        self.day = day if day is not None else self.df['Date'][0]
        self.dateTitle = datetime.datetime.strptime(self.day, "%Y-%m-%d").strftime("%b %d")
        self.dayFormat = datetime.datetime.strptime(self.day, "%Y-%m-%d").strftime("%m-%d-%Y")
        if field == 'WT' or field == 'NM':self.df = self.df.drop(['CP'], axis=1)
        if field == 'ET' or field == 'WB': 
            with open('data/prod/WB/wells.json', 'r') as f: wells:dict = json.loads(f.read())
            ww = [w.lower() for w in wells.keys()]
            mask = self.df['Well Name'].str.lower().isin(ww)
            if self.field == 'ET': mask = ~mask
            self.df = self.df[mask]
    
    def prepdf(self, df:pd.DataFrame):
        df = df[df['Well Name'] != 'South Texas Total']
        mon = self.day[:-3]
        df['Date'] = pd.to_datetime(df['Date'])
        df_mon = pd.DataFrame(df[df['Date'].dt.strftime('%Y-%m') == mon])
        ds = [col for col in df.columns if col in ['Date','TP','CP','Comments']]
        df_mon = df_mon.drop(ds,axis=1)
        df_monCuml = df_mon.groupby('Well Name').sum()
        df_monCuml = df_monCuml.rename({'Oil (BBLS)': 'MTD Oil','Gas (MCF)': 'MTD Gas','Water (BBLS)': 'MTD Water' },axis=1)

        df = pd.merge(df,df_monCuml,on='Well Name')
        if self.field == 'WT' or self.field == 'NM': col_ord = ['Well Name','Date','Oil (BBLS)','Gas (MCF)','Water (BBLS)', 'MTD Oil', 'MTD Gas', 'MTD Water','TP','Comments']
        else: col_ord = ['Well Name','Date','Oil (BBLS)','Gas (MCF)','Water (BBLS)', 'MTD Oil', 'MTD Gas', 'MTD Water','TP','CP','Comments']
        df = df.reindex(columns=col_ord)

        if self.wells is None:
            df = df[df['Date'] == self.day]
            df = df.drop(['Date'],axis=1).fillna('')
        else: 
            df = df.loc[df['Well Name'].isin(self.wells)]
            df = df.sort_values(['Date', 'Well Name'], ascending = [True , True])
            df['dateidx'] = pd.to_datetime(df['Date'])
            df.set_index('dateidx', inplace=True)
            df['MTD Oil'] = df.groupby(df.index.to_period('M'))['Oil (BBLS)'].cumsum()
            df['MTD Gas'] = df.groupby(df.index.to_period('M'))['Gas (MCF)'].cumsum()
            df['MTD Water'] = df.groupby(df.index.to_period('M'))['Water (BBLS)'].cumsum()

            df['Date'] = df['Date'].dt.date

            df = df.sort_values(['Date', 'Well Name'], ascending = [False , True])
        
        df = df._append(df[['Oil (BBLS)', 'Gas (MCF)', 'Water (BBLS)']].sum(),ignore_index=True)
        df.iloc[len(df)-1,df.columns.get_loc('Well Name')] = 'Total'

        for col in df.columns:
            if df[col].dtype == float:
                df[col] = np.round(df[col]).astype('Int64')
        widths = self.workoutWidth(df)
        df = df.reset_index(drop=True)
        

        df[df.select_dtypes(include='Int64').columns] = df.select_dtypes(include='Int64').astype('float64')
        df = df.replace(pd.NA,np.nan)
        return pd.DataFrame(df.fillna(' ')),widths
    
    def clean_text(self,text):
        return unicodedata.normalize('NFKD', text).encode('utf-8', 'ignore').decode('utf-8')

    def workoutWidth(self, df:pd.DataFrame):
        pdf = FPDF()
        pdf.set_font("Arial", size=7)
        widths = {}
        pos_end = pdf.l_margin
        for col in df.columns:
            widths[col] = {}
            longest = max(df[col].astype(str).tolist(),key=pdf.get_string_width)
    
            title = col.split(' ')
            if len(title) == 2:
                unit = '(MCF)' if title[1] == 'Gas' else '(BBLS)'
                if title[0] == 'Well': widths[col]['r1'] = col; widths[col]['r2'] = ''
                elif title[0] == 'MTD': widths[col]['r1'] = title[1]; widths[col]['r2'] = unit
                else: widths[col]['r1'] = title[0]; widths[col]['r2'] = title[1]
            else: widths[col]['r1'] = col; widths[col]['r2'] = ''
            
            offset = 6
            if col == 'TP' or col == 'CP': offset = 2

            widths[col]['width'] = pdf.get_string_width(longest) + offset
            pos_end += widths[col]['width']

            widths[col]['pos_end'] = pos_end
        pw = pdf.w - 2 * pdf.l_margin
        gifts = (pw - 30) - widths['Comments']['pos_end']
        if gifts > 0:
            gift = gifts/(len(widths)-1)
            for idx,poor in enumerate(widths):
                if widths[poor]['r1'] == 'Comments':continue
                widths[poor]['width'] += gift
                widths[poor]['pos_end'] += gift*(idx + 1)
        return widths
    
    def genReport(self):
        df,col_widths = self.prepdf(self.df)

        df[df.select_dtypes(include='float64').columns] = df.select_dtypes(include='float64').astype('Int64')
        pdf = FPDF()
        pdf.set_font("Arial", size=10)

        row_height = pdf.font_size * 1.2
        master_h = pdf.font_size * 1.2
        page_width = pdf.w - 2 * pdf.l_margin

        pdf.set_font("Arial", size=6)

        comm_offset = 0
        def header(first=False):
            pdf.set_font("Arial", style="B", size=16)
            pdf.cell(0, 10, txt=self.title, ln=True, align="C")
            if first: pdf.cell(0, 10, txt="CML Exploration LLC", ln=True, align="C")
            pdf.cell(0, 10, txt="Daily Gross Production Report", border='B', ln=True, align="C")
            pdf.set_font("Arial", size=10)
            pdf.cell(0, 8, txt=self.dayFormat, ln=True, align="R")

            pdf.set_x(col_widths['MTD Oil']['pos_end'])
            pdf.set_font("Arial", size=7)
            pdf.cell(0, 5, txt='**MTD**',ln=True)
            pdf.ln(0)

            for r in ['r1','r2']:
                brdr = 0
                if r == 'r2': pdf.cell(0, row_height, ln=True);pdf.set_font("Arial", size=6); brdr = 'B'
                for _,info in col_widths.items():
                    pdf.cell(0, row_height, txt=info[r],border=brdr)
                    pdf.set_x(info['pos_end'])
            pdf.cell(0, 5, txt='',ln=True)
            
            pdf.ln(row_height)
            nonlocal comm_offset
            comm_offset = col_widths['Comments']['pos_end'] - col_widths['Comments']['width']

        
        pdf.add_page()
        header(first=True)
        mb = pdf.b_margin
        k = 'TP' if self.field == 'NM' or self.field =='WT' else 'CP'
        for _, row in df.iterrows():
            if pdf.y + pdf.font_size * 1.5 > pdf.h - pdf.b_margin:
                pdf.add_page()
                header()
            pdf.set_font("Arial", size=6)

            cell_width = page_width - col_widths[k]['pos_end'] + pdf.l_margin
            if pdf.get_string_width(row['Comments']) + 2 > cell_width: row_height *= 2
            else: row_height = master_h
            
            for col in df.columns:
                if col == 'dbl':continue
                cell_value = str(row[col]).encode("latin-1", "ignore").decode("latin-1")
                cell_width = col_widths[col]['width']
                if col == 'Comments':
                    cell_width = page_width - col_widths[k]['pos_end'] + pdf.l_margin
                    comm_width = pdf.get_string_width(cell_value)
                    words = cell_value.split(' ')
                    row_height = master_h
                    if comm_width + 2 > cell_width: 
                        w = words[:]
                        cnt = 0
                        while comm_width + 2 > cell_width:
                            cnt += 1
                            w.pop()
                            comm1 = ' '.join(w)
                            comm_width = pdf.get_string_width(comm1)
                        
                        comm2 = ' '.join(words[-cnt:])
                        comm2_width = pdf.get_string_width(comm2)
                        
                        
                        pdf.cell(cell_width, row_height, txt=comm1, border='TR', ln=True, align='L')
                        if pdf.y + row_height > pdf.h - pdf.b_margin: pdf.set_auto_page_break(0,0)

                        fsize = 5.9
                        while comm2_width + 2 > cell_width:
                            pdf.set_font("Arial", size=fsize)
                            fsize -= .1
                            comm2_width = pdf.get_string_width(comm2)

                        pdf.cell(cell_width, row_height, txt='', border='L', ln=False, align='L')
                        pdf.set_x(comm_offset)
                        pdf.cell(cell_width, row_height, txt=comm2, border='RB', ln=False, align='L')

                        pdf.set_auto_page_break(1,mb)
                        continue   

                    
                pdf.cell(cell_width, row_height, txt=cell_value, border=1, ln=False, align='L')
            pdf.ln(row_height)


        pathtar = self.clientName if self.clientName else self.field
        reportTitle = f"{self.title} - {self.dateTitle}"
        folder_path = 'data/prod/reports/dailyreports/'

        files_to_rm = glob.glob(os.path.join(folder_path, f"*{self.title}*"))
        for file in files_to_rm:
            try:
                os.remove(file)
                print(f"Deleted: {file}")
            except Exception as e:
                print(f"Error deleting {file}: {e}")

        pdf_path = f'{folder_path+reportTitle}.pdf'
        csv_path = f'{folder_path+reportTitle}.csv'
        df.to_csv(csv_path,index=False)
        pdf.output(pdf_path,'F')  

        return {'pdf':pdf_path,'csv':csv_path,'title':reportTitle}
    
def totals_df(date=None):
    fields = ['ST','ET','WB','WT','NM','GC']
    area_map = {'ST':'SOUTH TEXAS','ET':'EAST TEXAS','WB':'EAST TEXAS WB',
                'WT':'WEST TEXAS','NM':'NEW MEXICO','GC':'GULF COAST'}
    date = dt.datetime.fromtimestamp(int(time.time()) - 60*60*24).strftime('%Y-%m-%d') if date is None else date
    res = {"Area": [], "BBL": [], 'MCF': []}
    oil_tot = 0; gas_tot = 0
    for field in fields:
        df = pd.read_csv(f'data/prod/{field}/data.csv')
        df = df[df['Date'] == date]
        if field == 'ET' or field == 'WB': 
            with open('data/prod/WB/wells.json', 'r') as f: wells:dict = json.loads(f.read())
            ww = [w.lower() for w in wells.keys()]
            mask = df['Well Name'].str.lower().isin(ww)
            if field == 'ET': mask = ~mask
            df = df[mask]
        oil = round(df['Oil (BBLS)'].sum(),0)
        gas = round(df['Gas (MCF)'].sum(),0)
        oil_tot += oil; gas_tot += gas
        res['Area'].append(area_map[field])
        res['BBL'].append(oil); res['MCF'].append(gas)
    res['Area'].append('Total')
    res['BBL'].append(round(oil_tot,0)); res['MCF'].append(round(gas_tot,0))

    df = pd.DataFrame(res)
    df['BBL'] = df['BBL'].astype(int)
    df['MCF'] = df['MCF'].astype(int)
    return df
    
def make_bold_underline(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.underline = True

def remove_table_borders(tbl):
    tbl_pr = tbl._element.find(qn('w:tblPr'))  # Get <w:tblPr>
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl._element.insert(0, tbl_pr)  # Insert at the beginning

    # Create or replace <w:tblBorders> to remove borders
    tbl_borders = tbl_pr.find(qn('w:tblBorders'))
    if tbl_borders is None:
        tbl_borders = OxmlElement('w:tblBorders')
        tbl_pr.append(tbl_borders)

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')  # 'nil' removes the border
        tbl_borders.append(border)

def cml_daily_docx():
    df = totals_df()    
    doc = Document()

    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for col_idx, column_name in enumerate(df.columns):
        hdr_cells[col_idx].text = column_name
        make_bold_underline(hdr_cells[col_idx])  # Apply bold and underline

    # Populate table with data
    rows = []
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for col_idx, value in enumerate(row):
            row_cells[col_idx].text = str(value)
        rows.append(row_cells)

    # Make last row bold & underlined
    last_row_cells = rows[-1]  # Get the last row added
    for cell in last_row_cells:
        make_bold_underline(cell)

    # Remove table borders
    remove_table_borders(table)

    doc.save('data/prod/reports/dailyreports/report.docx')
    return str(doc)


if __name__ == '__main__':
    reprt = ProdReport(field='ST',title='South Texas')
    for field,title in {'ST':'South Texas','ET':'East Texas','WT':'West TX','GC':'Gulf Coast','NM':'New Mexico'}.items():
        reprt = ProdReport(field=field,title=title)
        info = reprt.genReport()
    reprt = ProdReport(field='WB',title='Woodbine')
    reprt.genReport()
    cml_daily_docx()